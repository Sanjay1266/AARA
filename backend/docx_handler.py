import re
from typing import Dict, List
from docx import Document
from docx.shared import Pt

from backend.citation_engine import CitationEngine
from backend.bibliography_builder import BibliographyBuilder


class DocxHandler:
    """
    Handles all DOCX operations:
    - Reading paragraphs
    - Inserting citation markers
    - Finalizing document with real citations + bibliography
    """

    CITE_PATTERN = re.compile(r"\[CITE:\s*(.*?)\s*\|\s*(.*?)\]")

    # -----------------------------
    # READ
    # -----------------------------
    def read_paragraphs(self, docx_path: str) -> List[str]:
        document = Document(docx_path)
        return [p.text.strip() for p in document.paragraphs if p.text.strip()]

    # -----------------------------
    # WRITE (Markers)
    # -----------------------------
    def insert_citation_markers(
        self,
        input_docx: str,
        output_docx: str,
        citation_decisions: Dict[int, Dict]
    ):
        document = Document(input_docx)

        for idx, paragraph in enumerate(document.paragraphs):
            if idx not in citation_decisions:
                continue

            decision = citation_decisions[idx]
            if not decision.get("citation_required"):
                continue

            marker = f" [CITE: {decision['reference_id']} | {decision['confidence_score']}]"
            paragraph.add_run(marker)

        document.save(output_docx)

    # -----------------------------
    # FINALIZE
    # -----------------------------
    def finalize_document(
        self,
        input_docx: str,
        output_docx: str,
        reference_metadata: Dict[str, Dict],
        citation_style: str = "APA"
    ):
        document = Document(input_docx)

        citation_engine = CitationEngine(style=citation_style)
        bibliography_builder = BibliographyBuilder(style=citation_style)

        used_refs = []

        for paragraph in document.paragraphs:
            matches = list(self.CITE_PATTERN.finditer(paragraph.text))
            if not matches:
                continue

            new_text = paragraph.text

            for match in matches:
                ref_id = match.group(1).strip()
                if ref_id not in reference_metadata:
                    continue

                citation = citation_engine.format_in_text(
                    ref_id, reference_metadata[ref_id]
                )

                new_text = new_text.replace(match.group(0), citation)

                if ref_id not in used_refs:
                    used_refs.append(ref_id)

            paragraph.clear()
            run = paragraph.add_run(new_text)
            run.font.size = Pt(11)

        self._append_bibliography(
            document, used_refs, reference_metadata, bibliography_builder
        )

        document.save(output_docx)

    # -----------------------------
    # HELPERS
    # -----------------------------
    def _append_bibliography(
        self,
        document: Document,
        used_refs,
        reference_metadata,
        bibliography_builder
    ):
        document.add_page_break()

        # SAFE heading insertion
        try:
            document.add_heading("References", level=1)
        except KeyError:
            p = document.add_paragraph("References")
            p.runs[0].bold = True

        for i, ref_id in enumerate(used_refs, start=1):
            meta = reference_metadata[ref_id].copy()
            meta["index"] = i

            entry = bibliography_builder.build_entry(ref_id, meta)
            p = document.add_paragraph(entry)
            p.paragraph_format.space_after = Pt(6)
